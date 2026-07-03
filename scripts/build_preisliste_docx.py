from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables"
OUTPUT_DIR.mkdir(exist_ok=True)

DOCX_PATH = OUTPUT_DIR / "Blitzkneisser_Preisliste_2026_27.docx"

ACCENT = RGBColor(0xAA, 0x87, 0x59)
INK = RGBColor(0x1D, 0x1D, 0x1D)
MUTED = RGBColor(0x6F, 0x6A, 0x66)
LINE = "D9D1C7"
SOFT = "F6F1EB"

IMAGES = {
    "cover": ROOT / "assets/uploads/Blitzkneisser-Elopement-Cadini-Seceda-14-home.jpg",
    "story": ROOT / "assets/uploads/Blitzkneisser-Proposal-Dolomites-177-home.jpg",
    "collections": ROOT / "assets/uploads/Blitzkneisser-Dolomites-Elopement-Wedding-154.jpg",
    "formats": ROOT / "assets/uploads/Blitzkneisser-Dolomites-Elopement-INSTA-7-home.jpg",
    "details": ROOT / "assets/uploads/Blitzkneisser-Pizza-Elopement-Dolomits-2-home.jpg",
    "closing": ROOT / "assets/uploads/Blitzkneisser-Dolomites-Elopement-Wedding-236.jpg",
}

CONTENT = {
    "eyebrow": "INVESTMENT GUIDE 2026 / 2027",
    "title": "ZEITLOSE ERINNERUNGEN",
    "subtitle": (
        "Hochzeitsfotografie für Tirol, Innsbruck und die Dolomiten.\n"
        "Ruhig, authentisch und mit einem klaren Gefühl für euren Tag."
    ),
    "intro": (
        "Zwischen Bergen, Licht und Weite entstehen Hochzeitsbilder, die ruhig wirken, "
        "authentisch bleiben und eure Geschichte einfühlsam erzählen. Diese Preisliste "
        "gibt euch eine klare Orientierung und zeigt, welche Begleitungen für euren Tag möglich sind."
    ),
    "story_title": "WAS IHR VON MIR ERWARTEN KÖNNT",
    "story_copy": [
        (
            "Die Tiroler Alpen sind für mich einer der eindrucksvollsten Orte für ehrliche "
            "und zeitlose Hochzeitsbilder. Mein Blick bleibt dabei immer auf euch gerichtet: "
            "auf eure Verbindung, eure Ruhe und die Momente dazwischen."
        ),
        (
            "Ich lebe und arbeite in Innsbruck und kenne viele besondere Orte in Tirol und "
            "den Alpen, an denen intime Hochzeiten, Elopements und ruhige Paarmomente ihre "
            "eigene Magie entfalten. Von der ersten Idee bis zur Begleitung vor Ort entsteht "
            "so ein Ablauf, der sich leicht und stimmig anfühlt."
        ),
        (
            "Wenn ihr eure Hochzeit oder euer Elopement in Tirol plant, begleite ich euch "
            "vom ersten Gedanken bis hin zu den leisen und großen Momenten eures Tages. "
            "So entstehen Bilder, die nicht nur Erinnerungen bewahren, sondern eure Geschichte "
            "auch Jahre später noch fühlbar machen."
        ),
    ],
    "included_title": "IN JEDER BEGLEITUNG ENTHALTEN",
    "included": [
        "persönliche Abstimmung zu Ablauf, Licht und Stimmung",
        "ruhige fotografische Begleitung mit Fokus auf echte Momente",
        "sorgfältige Bildauswahl und hochwertige finale Bearbeitung",
        "digitale Auslieferung eurer finalen Bilder",
        "ehrliche Beratung, welche Begleitung wirklich zu eurem Tag passt",
    ],
    "collections_title": "HOCHZEITSREPORTAGEN",
    "collections_intro": (
        "Die beiden Hauptbegleitungen sind bewusst klar gehalten und dienen als Orientierung. "
        "Jedes Angebot kann an eure Pläne, eure Gästezahl und den Charakter eures Tages angepasst werden."
    ),
    "collections_note": (
        "Bei Hochzeiten mit mehr als 100 Gästen empfehle ich einen zweiten Fotografen, "
        "damit kein wichtiger Moment verloren geht."
    ),
    "formats_title": "ELOPEMENTS & KLEINERE FORMATE",
    "formats_intro": (
        "Nicht jede Geschichte braucht einen ganzen Hochzeitstag. Für intime Hochzeiten, "
        "kleine standesamtliche Feiern oder ergänzende Kapitel rund um eure Hochzeit "
        "gibt es bewusst reduzierte Formate."
    ),
    "extras_title": "ZUSATZOPTIONEN & WEITERE KAPITEL",
    "extras_intro": (
        "Wenn euer Tag mehr Zeit, zusätzliche Erlebnisse oder bewegte Erinnerungen braucht, "
        "können wir die Begleitung gezielt erweitern."
    ),
    "process_title": "SO GEHT ES WEITER",
    "process": [
        ("01", "Anfrage", "Ihr schreibt mir mit euren Eckdaten und eurer Idee."),
        ("02", "Abstimmung", "Wir besprechen den Ablauf, die Stimmung und die passende Begleitung."),
        ("03", "Euer Tag", "Ich begleite euch ruhig, unaufdringlich und mit einem klaren Blick für echte Momente."),
        ("04", "Erinnerungen", "Ihr erhaltet eure finalen Bilder sorgfältig kuratiert und hochwertig bearbeitet."),
    ],
    "closing_title": "LASST UNS EURE GESCHICHTE PLANEN",
    "closing_copy": (
        "Wenn ihr euch eine Hochzeitsreportage wünscht, die sich leicht anfühlt und "
        "eure Verbindung ehrlich erzählt, freue ich mich sehr, mehr über euch zu erfahren. "
        "Schreibt mir einfach mit eurem Wunschdatum und euren ersten Ideen."
    ),
    "contact_lines": [
        "Andreas Kiss · Blitzkneisser Fotografie",
        "Innsbruck, Tirol",
        "foto@blitzkneisser.com",
        "www.hochzeitsfotograf.tirol",
    ],
}

PRICES = {
    "package_1": {
        "title": "REPORTAGE · BIS ZU 8 STUNDEN",
        "price": "4.200 €",
        "items": [
            "durchgehende fotografische Begleitung mit einem Fotografen",
            "zum Beispiel Getting Ready, Trauung, Paarfotos und Abendmomente",
            "mindestens 500 bearbeitete Bilder",
            "ideal für Hochzeiten mit klarem Ablauf und vielen wichtigen Programmpunkten",
        ],
    },
    "package_2": {
        "title": "REPORTAGE · BIS ZU 6 STUNDEN",
        "price": "3.700 €",
        "items": [
            "durchgehende fotografische Begleitung mit einem Fotografen",
            "zum Beispiel Trauung, Paarfotos und wesentliche Teile des Tages",
            "mindestens 400 bearbeitete Bilder",
            "ideal für intime Feiern mit fokussierter Begleitung",
        ],
    },
    "elopement": {
        "title": "ELOPEMENT",
        "price": "ab 3.500 €",
        "copy": "Für Paare, die ihren Tag klein, bewusst und in den Bergen erleben möchten.",
    },
    "civil": {
        "title": "STANDESAMT",
        "price": "1.200 €",
        "copy": "2 Stunden fotografische Begleitung an Wochentagen, inklusive ca. 100 bearbeiteten Bildern.",
    },
    "after": {
        "title": "PRE / AFTER WEDDING",
        "price": "1.500 €",
        "copy": "Ein eigenes Kapitel für euch zwei, wenn am Hochzeitstag mehr Raum für gemeinsame Bilder entstehen soll.",
    },
    "film": {
        "title": "FILM",
        "price": "ab 3.800 €",
        "copy": "Bewegte Erinnerungen als ergänzender Hochzeitsfilm, passend zu den Bildern und dem Charakter eures Tages.",
    },
    "extra_hour": {
        "title": "ZUSATZSTUNDE",
        "price": "300 €",
        "copy": "Wenn euer Tag länger wird oder wir spontan mehr Zeit brauchen.",
    },
    "photobox": {
        "title": "FOTOBOX",
        "price": "800 €",
        "copy": "Für Feiern, bei denen eure Gäste ein spielerisches, direktes Andenken mitnehmen sollen.",
    },
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:" + key), str(value))


def set_table_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name, size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = color


def add_text(paragraph, text, *, font="Arial", size=11, bold=False, color=INK, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, font, size=size, bold=bold, color=color, italic=italic)
    return run


def configure_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)
    return doc


def define_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = INK

    doc.styles["Heading 1"].font.size = Pt(20)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)

    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 2"].font.color.rgb = ACCENT
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(6)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(4)

    doc.styles["Heading 3"].font.size = Pt(11)
    doc.styles["Heading 3"].font.bold = True
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(4)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(2)

    doc.styles["List Bullet"].font.name = "Arial"
    doc.styles["List Bullet"].font.size = Pt(11)
    doc.styles["List Bullet"].paragraph_format.space_after = Pt(4)
    doc.styles["List Bullet"].paragraph_format.line_spacing = 1.1


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Blitzkneisser Fotografie  ·  foto@blitzkneisser.com  ·  hochzeitsfotograf.tirol")
    set_run_font(run, "Arial", size=8, color=MUTED)


def add_eyebrow(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    set_run_font(run, "Arial", size=9, bold=True, color=ACCENT)
    run.font.all_caps = True
    paragraph.paragraph_format.keep_with_next = True


def add_display_title(doc, text):
    for line in text.split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(line)
        set_run_font(run, "Arial", size=28, bold=True, color=INK)


def add_intro_block(doc):
    add_eyebrow(doc, CONTENT["eyebrow"])
    add_display_title(doc, CONTENT["title"])
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(10)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.paragraph_format.keep_with_next = True
    for i, line in enumerate(CONTENT["subtitle"].split("\n")):
        if i:
            subtitle.add_run().add_break()
        add_text(subtitle, line, font="Arial", size=12, color=MUTED)

    intro = doc.add_paragraph(CONTENT["intro"])
    intro.paragraph_format.space_after = Pt(14)
    intro.paragraph_format.keep_with_next = True

    doc.add_picture(str(IMAGES["cover"]), width=Inches(6.65))


def add_section_title(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.keep_with_next = True
    add_text(p, title, font="Arial", size=20, bold=True, color=INK)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p_pr = p._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_story_page(doc):
    doc.add_page_break()
    add_section_title(doc, CONTENT["story_title"])
    for paragraph in CONTENT["story_copy"]:
        doc.add_paragraph(paragraph)
    add_divider(doc)
    add_section_title(doc, CONTENT["included_title"])
    for item in CONTENT["included"]:
        p = doc.add_paragraph(style="List Bullet")
        add_text(p, item, font="Arial", size=11, color=INK)
    doc.add_paragraph()
    doc.add_picture(str(IMAGES["story"]), width=Inches(6.65))


def style_package_cell(cell, shaded=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_cell_margins(cell, top=120, start=140, bottom=120, end=140)
    set_cell_border(
        cell,
        top={"val": "single", "sz": 8, "color": LINE},
        bottom={"val": "single", "sz": 8, "color": LINE},
        left={"val": "single", "sz": 8, "color": LINE},
        right={"val": "single", "sz": 8, "color": LINE},
    )
    if shaded:
        set_cell_shading(cell, SOFT)


def add_package_box(doc, title, price, items):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)
    left, right = table.rows[0].cells
    style_package_cell(left)
    style_package_cell(right, shaded=True)

    p = left.paragraphs[0]
    p.paragraph_format.space_after = Pt(8)
    add_text(p, title, font="Arial", size=14, bold=True, color=INK)
    for item in items:
        bullet = left.add_paragraph(style="List Bullet")
        add_text(bullet, item, font="Arial", size=10.5, color=INK)

    price_paragraph = right.paragraphs[0]
    price_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    price_paragraph.paragraph_format.space_after = Pt(8)
    add_text(price_paragraph, "ab", font="Arial", size=10, bold=False, color=MUTED)
    price_paragraph.add_run().add_break()
    add_text(price_paragraph, price, font="Arial", size=18, bold=True, color=ACCENT)
    doc.add_paragraph()


def add_collections_page(doc):
    doc.add_page_break()
    add_section_title(doc, CONTENT["collections_title"])
    doc.add_paragraph(CONTENT["collections_intro"])
    add_package_box(
        doc,
        PRICES["package_1"]["title"],
        PRICES["package_1"]["price"],
        PRICES["package_1"]["items"],
    )
    add_package_box(
        doc,
        PRICES["package_2"]["title"],
        PRICES["package_2"]["price"],
        PRICES["package_2"]["items"],
    )
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(10)
    add_text(note, CONTENT["collections_note"], font="Arial", size=10.5, italic=True, color=MUTED)
    doc.add_picture(str(IMAGES["collections"]), width=Inches(6.65))


def add_format_row(table, row_idx, title, price, copy):
    left, right = table.rows[row_idx].cells
    style_package_cell(left)
    style_package_cell(right, shaded=True)
    lp = left.paragraphs[0]
    lp.paragraph_format.space_after = Pt(4)
    add_text(lp, title, font="Arial", size=12.5, bold=True, color=INK)
    cp = left.add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    add_text(cp, copy, font="Arial", size=10.5, color=INK)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(rp, price, font="Arial", size=15, bold=True, color=ACCENT)


def add_formats_page(doc):
    doc.add_page_break()
    add_section_title(doc, CONTENT["formats_title"])
    doc.add_paragraph(CONTENT["formats_intro"])
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)
    add_format_row(table, 0, PRICES["elopement"]["title"], PRICES["elopement"]["price"], PRICES["elopement"]["copy"])
    add_format_row(table, 1, PRICES["civil"]["title"], PRICES["civil"]["price"], PRICES["civil"]["copy"])
    add_format_row(table, 2, PRICES["after"]["title"], PRICES["after"]["price"], PRICES["after"]["copy"])
    add_format_row(table, 3, PRICES["film"]["title"], PRICES["film"]["price"], PRICES["film"]["copy"])
    doc.add_paragraph()
    doc.add_picture(str(IMAGES["formats"]), width=Inches(6.65))


def add_simple_offer_table(doc, offers):
    table = doc.add_table(rows=len(offers), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(4.8)
    table.columns[1].width = Inches(1.7)
    for idx, (title, price, copy) in enumerate(offers):
        left, right = table.rows[idx].cells
        style_package_cell(left)
        style_package_cell(right, shaded=True)
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        add_text(p, title, font="Arial", size=12.5, bold=True, color=INK)
        body = left.add_paragraph()
        body.paragraph_format.space_after = Pt(0)
        add_text(body, copy, font="Arial", size=10.5, color=INK)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(rp, price, font="Arial", size=14, bold=True, color=ACCENT)


def add_extras_page(doc):
    doc.add_page_break()
    add_section_title(doc, CONTENT["extras_title"])
    doc.add_paragraph(CONTENT["extras_intro"])
    add_simple_offer_table(
        doc,
        [
            (PRICES["extra_hour"]["title"], PRICES["extra_hour"]["price"], PRICES["extra_hour"]["copy"]),
            (PRICES["photobox"]["title"], PRICES["photobox"]["price"], PRICES["photobox"]["copy"]),
        ],
    )
    doc.add_paragraph()
    doc.add_picture(str(IMAGES["details"]), width=Inches(6.65))


def add_process_page(doc):
    doc.add_page_break()
    add_section_title(doc, CONTENT["process_title"])
    for step_no, title, copy in CONTENT["process"]:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(4)
        heading.paragraph_format.space_after = Pt(1)
        add_text(heading, f"{step_no}  ", font="Arial", size=10, bold=True, color=ACCENT)
        add_text(heading, title.upper(), font="Arial", size=11, bold=True, color=INK)
        body = doc.add_paragraph(copy)
        body.paragraph_format.left_indent = Inches(0.32)
        body.paragraph_format.space_after = Pt(8)

    add_divider(doc)
    add_section_title(doc, CONTENT["closing_title"])
    doc.add_paragraph(CONTENT["closing_copy"])
    contact = doc.add_paragraph()
    contact.paragraph_format.space_before = Pt(8)
    for idx, line in enumerate(CONTENT["contact_lines"]):
        if idx:
            contact.add_run().add_break()
        add_text(contact, line, font="Arial", size=11, bold=(idx == 0), color=INK if idx == 0 else MUTED)
    doc.add_paragraph()
    doc.add_picture(str(IMAGES["closing"]), width=Inches(6.65))


def main():
    doc = configure_document()
    define_styles(doc)
    add_footer(doc.sections[0])
    add_intro_block(doc)
    add_story_page(doc)
    add_collections_page(doc)
    add_formats_page(doc)
    add_extras_page(doc)
    add_process_page(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
